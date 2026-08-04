"""Safe export rendering for final lecture notes."""
from __future__ import annotations

import re
from io import BytesIO

from fastapi.responses import StreamingResponse


def safe_filename(value: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip(".-")
    return clean[:100] or "lecture-notes"


def _strip_math_delimiters(markdown: str) -> str:
    """PDF/HTML export can't run KaTeX, so present math as plain readable text
    by removing the `$`/`$$` delimiters rather than leaving them literal."""
    text = re.sub(r"\$\$(.+?)\$\$", r"\1", markdown, flags=re.DOTALL)
    text = re.sub(r"\$(.+?)\$", r"\1", text)
    return text


def _markdown_to_html(markdown: str) -> str:
    """Render Markdown to HTML. Uses python-markdown when available; falls back
    to a minimal renderer so export never dumps raw `#`/`*` markup."""
    source = _strip_math_delimiters(markdown or "")
    try:
        import markdown as md

        return md.markdown(source, extensions=["extra", "sane_lists"])
    except Exception:
        # Minimal fallback: headings, bullets, and paragraphs.
        lines_out = []
        in_list = False
        for raw in source.splitlines():
            line = raw.rstrip()
            heading = re.match(r"^(#{1,6})\s+(.*)$", line)
            bullet = re.match(r"^[*-]\s+(.*)$", line)
            if heading:
                if in_list:
                    lines_out.append("</ul>")
                    in_list = False
                level = len(heading.group(1))
                lines_out.append(f"<h{level}>{heading.group(2)}</h{level}>")
            elif bullet:
                if not in_list:
                    lines_out.append("<ul>")
                    in_list = True
                lines_out.append(f"<li>{bullet.group(1)}</li>")
            elif line:
                if in_list:
                    lines_out.append("</ul>")
                    in_list = False
                lines_out.append(f"<p>{line}</p>")
        if in_list:
            lines_out.append("</ul>")
        return "\n".join(lines_out)


_PDF_CSS = """
  body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #1f2937; line-height: 1.5;
         margin: 2.5cm 2cm; }
  h1 { font-size: 22pt; border-bottom: 2px solid #6366f1; padding-bottom: 6px; }
  h2 { font-size: 16pt; margin-top: 22px; color: #3730a3; }
  h3 { font-size: 13pt; margin-top: 16px; color: #4338ca; }
  ul { margin: 6px 0 6px 18px; } li { margin: 3px 0; }
  p { margin: 8px 0; }
  code { background: #f3f4f6; padding: 1px 4px; border-radius: 3px; }
"""


def export_notes(title: str, markdown: str, export_format: str) -> StreamingResponse:
    filename = safe_filename(title)

    if export_format == "md":
        return StreamingResponse(
            iter([markdown.encode()]),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}.md"'},
        )

    if export_format == "txt":
        text = _strip_math_delimiters(markdown)
        text = re.sub(r"[*_#>`]", "", text)
        return StreamingResponse(
            iter([text.encode()]),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{filename}.txt"'},
        )

    if export_format == "docx":
        from docx import Document

        document = Document()
        document.add_heading(title, 0)
        for line in _strip_math_delimiters(markdown).splitlines():
            if line.startswith("### "):
                document.add_heading(line[4:], 3)
            elif line.startswith("## "):
                document.add_heading(line[3:], 2)
            elif line.startswith("# "):
                document.add_heading(line[2:], 1)
            elif line.strip().startswith(("- ", "* ")):
                document.add_paragraph(line.strip()[2:], style="List Bullet")
            elif line.strip():
                document.add_paragraph(line)
        output = BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
        )

    if export_format == "pdf":
        try:
            from weasyprint import HTML
        except ImportError as exc:
            raise RuntimeError("PDF export is not installed") from exc
        body = _markdown_to_html(markdown)
        html = f"<!doctype html><html><head><meta charset='utf-8'><style>{_PDF_CSS}</style></head><body><h1>{title}</h1>{body}</body></html>"
        return StreamingResponse(
            iter([HTML(string=html).write_pdf()]),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
        )

    raise ValueError("Unsupported export format")
